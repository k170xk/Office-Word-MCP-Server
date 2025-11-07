"""
Document-level formatting tools for Word Document Server.
Handles default fonts, headers, footers, and document-wide settings.
"""
import os
from typing import Optional
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


async def set_default_font(filename: str, font_name: str = "Calibri", font_size: int = 11, apply_to_existing: bool = True) -> str:
    """
    Set the default font for the entire document.
    This affects the Normal style and all paragraphs that use it.
    
    Args:
        filename: Path to the Word document
        font_name: Font family name (default: Calibri)
        font_size: Font size in points (default: 11)
        apply_to_existing: If True, apply font to existing paragraphs (default: True)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    try:
        doc = Document(filename)
        
        # Get or create the Normal style
        try:
            normal_style = doc.styles['Normal']
        except KeyError:
            normal_style = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        
        # Set font properties for Normal style
        font = normal_style.font
        font.name = font_name
        font.size = Pt(font_size)
        
        # Also update all existing paragraphs that use Normal style
        if apply_to_existing:
            for paragraph in doc.paragraphs:
                if paragraph.style.name == 'Normal' or paragraph.style.name.startswith('Normal'):
                    for run in paragraph.runs:
                        if not run.font.name or run.font.name == 'None':
                            run.font.name = font_name
                        if not run.font.size or run.font.size is None:
                            run.font.size = Pt(font_size)
        
        # Save the document
        doc.save(filename)
        
        return f"Default font set to {font_name} {font_size}pt for document {filename}"
    except Exception as e:
        return f"Failed to set default font: {str(e)}"


async def update_header_title_subtitle(filename: str, title: Optional[str] = None, subtitle: Optional[str] = None) -> str:
    """
    Update the header title and subtitle in a Word document.
    This modifies the document header section.
    
    Args:
        filename: Path to the Word document
        title: Header title text (optional, None to leave unchanged)
        subtitle: Header subtitle text (optional, None to leave unchanged)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    
    if title is None and subtitle is None:
        return "At least one of 'title' or 'subtitle' must be provided"
    
    try:
        doc = Document(filename)
        
        # Get the header section (first section's header)
        if len(doc.sections) == 0:
            return "Document has no sections"
        
        header = doc.sections[0].header
        
        # Clear existing header paragraphs if we're updating
        if title is not None or subtitle is not None:
            # Clear all existing paragraphs in header
            for paragraph in header.paragraphs:
                paragraph.clear()
        
        # Add title if provided
        if title is not None:
            title_para = header.paragraphs[0] if len(header.paragraphs) > 0 else header.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = "Calibri"
            title_run.font.size = Pt(11)
            title_run.bold = True
            title_para.alignment = 1  # Center alignment
        
        # Add subtitle if provided
        if subtitle is not None:
            subtitle_para = header.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.name = "Calibri"
            subtitle_run.font.size = Pt(11)
            subtitle_para.alignment = 1  # Center alignment
        
        # Save the document
        doc.save(filename)
        
        result_parts = []
        if title is not None:
            result_parts.append(f"title: '{title}'")
        if subtitle is not None:
            result_parts.append(f"subtitle: '{subtitle}'")
        
        return f"Header updated with {', '.join(result_parts)} in document {filename}"
    except Exception as e:
        return f"Failed to update header: {str(e)}"


async def get_header_info(filename: str) -> str:
    """
    Get information about the current header content.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        
        if len(doc.sections) == 0:
            return "Document has no sections"
        
        header = doc.sections[0].header
        
        if len(header.paragraphs) == 0:
            return "Header is empty"
        
        info_parts = []
        for i, para in enumerate(header.paragraphs):
            text = para.text.strip()
            if text:
                info_parts.append(f"Line {i+1}: {text}")
        
        if not info_parts:
            return "Header exists but contains no text"
        
        return "Header content:\n" + "\n".join(info_parts)
    except Exception as e:
        return f"Failed to get header info: {str(e)}"

